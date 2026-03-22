from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
import os

def convert_pdf_to_word(pdf_path: str, docx_path: str):
    """
    Converts PDF to Word (DOCX).
    Attempts full layout reconstruction first, 
    falls back to clean text extraction if the layout is too complex or treated as an image.
    """
    try:
        if not os.path.exists(pdf_path):
            print(f"Error: File not found {pdf_path}")
            return False
            
        # Attempt 1: Standard pdf2docx conversion (Layout Optimized)
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None, multi_processing=True)
            cv.close()
            
            # Check if the output file is too small or if we can do better
            if os.path.exists(docx_path) and os.path.getsize(docx_path) > 1000:
                return True
        except Exception as e:
            print(f"Standard Conversion failed, falling back: {e}")

        # Attempt 2: Clean Text Extraction (Text Optimized)
        # This ensures that text is extracted as actual text (not images)
        return convert_to_word_clean_text(pdf_path, docx_path)

    except Exception as e:
        print(f"PDF to Word Conversion Error: {e}")
        return False

def convert_to_word_clean_text(pdf_path: str, docx_path: str):
    """
    Extracts purely text from PDF using PyMuPDF and writes it into a clean Word document.
    Ensures all text is editable.
    """
    try:
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page in doc:
            # Extract text blocks
            text = page.get_text("text")
            if text.strip():
                word_doc.add_paragraph(text)
                word_doc.add_page_break()
        
        if len(word_doc.paragraphs) == 0:
            # If no text was found, it's likely a scan (image-only)
            # We can't do much without OCR, but we should notify or try one last thing
            word_doc.add_paragraph("No selectable text found in this PDF. It might be a scanned image.")
            
        word_doc.save(docx_path)
        doc.close()
        return True
    except Exception as e:
        print(f"Clean Text Conversion Error: {e}")
        return False
